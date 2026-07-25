import os
import json
import time
import uuid
from typing import List, Dict
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from groq import Groq
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = FastAPI(title="Autonomous Document Agent")

# ================== YOUR GROQ KEY ==================
GROQ_API_KEY = ""
# ===================================================

client = Groq(api_key=GROQ_API_KEY)

class AgentRequest(BaseModel):
    request: str

class Task(BaseModel):
    step: int
    description: str
    expected_output: str

class ExecutionPlan(BaseModel):
    document_type: str
    title: str
    tasks: List[Task]

def call_llm(prompt: str, temperature: float = 0.7, max_tokens: int = 2000) -> str:
    for attempt in range(3):
        try:
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=temperature,
                max_tokens=max_tokens,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            if attempt == 2:
                raise HTTPException(500, f"Groq Error: {str(e)}")
            time.sleep(2)
    return ""

def generate_plan(user_request: str) -> ExecutionPlan:
    prompt = f"""You are an expert autonomous business document agent.
User Request: {user_request}

Return ONLY valid JSON in this format:
{{
  "document_type": "Project Proposal / SOP / Business Report etc.",
  "title": "Clear Professional Title",
  "tasks": [
    {{"step": 1, "description": "Task description", "expected_output": "What to generate"}}
  ]
}}
Create 5-8 useful tasks. Use mock data if needed."""

    response_text = call_llm(prompt, temperature=0.3, max_tokens=1500)
    try:
        return ExecutionPlan(**json.loads(response_text))
    except:
        # Fallback
        return ExecutionPlan(
            document_type="Business Document",
            title="Autonomous Document",
            tasks=[Task(step=i, description=f"Section {i}", expected_output="Content") for i in range(1,6)]
        )

def execute_task(task: Task, context: str) -> str:
    prompt = f"""Context: {context}

Task {task.step}: {task.description}
Expected: {task.expected_output}

Write professional business content."""
    return call_llm(prompt, temperature=0.75)

def create_docx(plan: ExecutionPlan, sections: Dict[int, str], filename: str):
    doc = Document()
    doc.add_heading(plan.title, 0)
    doc.add_paragraph(f"Type: {plan.document_type} | Generated: {time.strftime('%Y-%m-%d')}")
    doc.add_page_break()

    for task in plan.tasks:
        content = sections.get(task.step, "No content generated.")
        doc.add_heading(f"{task.step}. {task.description}", level=1)
        for line in content.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph()

    doc.save(filename)

@app.post("/agent")
async def run_agent(req: AgentRequest):
    if len(req.request.strip()) < 10:
        raise HTTPException(400, "Request too short.")

    plan = generate_plan(req.request)
    
    sections = {}
    context = f"Request: {req.request}"

    for task in plan.tasks:
        sections[task.step] = execute_task(task, context)
        context += f"\n--- Section {task.step} done ---"

    filename = f"doc_{uuid.uuid4().hex[:10]}.docx"
    create_docx(plan, sections, filename)

    return FileResponse(
        path=filename,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.get("/")
async def home():
    return {"message": "Agent is ready! POST to /agent"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)