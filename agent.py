import os
import json
import time
import uuid
from typing import List, Dict

from groq import Groq
from pydantic import BaseModel
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ======================================================
# Load Environment Variables
# ======================================================

import os
import streamlit as st
from dotenv import load_dotenv
from groq import Groq

load_dotenv()

GROQ_API_KEY = None

# Streamlit Cloud
try:
    GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
except Exception:
    pass

# Docker / Render / Local env
if not GROQ_API_KEY:
    GROQ_API_KEY = os.getenv("GROQ_API_KEY")

if not GROQ_API_KEY:
    raise ValueError(
        "GROQ_API_KEY not found."
    )

client = Groq(api_key=GROQ_API_KEY)

# ======================================================
# Models
# ======================================================

class Task(BaseModel):
    step: int
    description: str
    expected_output: str


class ExecutionPlan(BaseModel):
    document_type: str
    title: str
    tasks: List[Task]


class AgentResult(BaseModel):
    filename: str
    title: str
    document_type: str
    sections: Dict[int, str]


# ======================================================
# LLM Wrapper
# ======================================================

def call_llm(
    prompt: str,
    temperature: float = 0.7,
    max_tokens: int = 2000,
    retries: int = 3
) -> str:

    last_error = None

    for attempt in range(retries):

        try:

            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an expert AI Document Agent. "
                            "Always generate professional, well-structured content."
                        ),
                    },
                    {
                        "role": "user",
                        "content": prompt,
                    },
                ],
                temperature=temperature,
                max_tokens=max_tokens,
            )

            return response.choices[0].message.content.strip()

        except Exception as e:

            last_error = e
            time.sleep(2)

    raise Exception(f"Groq API Error: {last_error}")


# ======================================================
# Utility
# ======================================================

def clean_json(text: str) -> str:
    """
    Removes markdown formatting if the model returns JSON
    inside ```json blocks.
    """

    text = text.replace("```json", "")
    text = text.replace("```", "")

    return text.strip()


def safe_filename(title: str) -> str:
    """
    Converts title into a safe filename.
    """

    invalid = '\\/:*?"<>|'

    for ch in invalid:
        title = title.replace(ch, "")

    title = title.replace(" ", "_")

    return title[:60]
# ======================================================
# Memory
# ======================================================

class AgentMemory:
    """
    Stores context generated during execution.
    """

    def __init__(self):
        self.history = []

    def add(self, title: str, content: str):
        self.history.append(
            {
                "title": title,
                "content": content
            }
        )

    def get_context(self):

        if not self.history:
            return "No previous sections generated."

        context = ""

        for item in self.history:

            context += f"""

Section:
{item['title']}

Content:
{item['content']}

"""

        return context


# ======================================================
# Fallback Plan
# ======================================================

def fallback_plan(user_request: str):

    return ExecutionPlan(
        document_type="Business Document",

        title="AI Generated Document",

        tasks=[

            Task(
                step=1,
                description="Executive Summary",
                expected_output="Overview"
            ),

            Task(
                step=2,
                description="Problem Statement",
                expected_output="Current problem"
            ),

            Task(
                step=3,
                description="Proposed Solution",
                expected_output="Detailed solution"
            ),

            Task(
                step=4,
                description="Implementation Plan",
                expected_output="Execution roadmap"
            ),

            Task(
                step=5,
                description="Conclusion",
                expected_output="Final summary"
            )

        ]
    )


# ======================================================
# Plan Generator
# ======================================================

def generate_plan(user_request: str):

    planning_prompt = f"""
You are an Autonomous AI Planning Agent.

Your task is NOT to write the document.

Instead,
break the request into logical execution steps.

User Request

{user_request}

Return ONLY valid JSON.

Format

{{
    "document_type":"Business Proposal",
    "title":"Professional Title",

    "tasks":[

        {{
            "step":1,
            "description":"Executive Summary",
            "expected_output":"Summary"
        }}

    ]
}}

Rules

Generate 5-8 tasks.

Tasks must be sequential.

Each task should describe ONE section.

No markdown.

No explanation.

Only JSON.
"""

    response = call_llm(
        planning_prompt,
        temperature=0.2,
        max_tokens=1400
    )

    response = clean_json(response)

    try:

        data = json.loads(response)

        return ExecutionPlan(**data)

    except Exception as e:

        print("Planner Error:", e)

        return fallback_plan(user_request)
    
# ======================================================
# Task Executor
# ======================================================

def execute_task(
    task: Task,
    user_request: str,
    memory: AgentMemory
) -> str:

    previous_context = memory.get_context()

    prompt = f"""
You are an expert business writer.

User Request:
{user_request}

Previously Generated Sections:

{previous_context}

--------------------------------------------------

Current Task

Step Number:
{task.step}

Section:
{task.description}

Expected Output:
{task.expected_output}

--------------------------------------------------

Instructions

Write ONLY this section.

Do NOT rewrite previous sections.

Write professionally.

Use headings, bullets and paragraphs where appropriate.

Do not include explanations outside the document.

Generate detailed content.
"""

    content = call_llm(
        prompt=prompt,
        temperature=0.7,
        max_tokens=1800
    )

    memory.add(task.description, content)

    return content


# ======================================================
# Execute Full Plan
# ======================================================

def execute_plan(
    plan: ExecutionPlan,
    user_request: str,
    progress_callback=None
):

    memory = AgentMemory()

    sections = {}

    total = len(plan.tasks)

    for index, task in enumerate(plan.tasks):

        if progress_callback:
            progress_callback(index, total, task.description)

        content = execute_task(
            task=task,
            user_request=user_request,
            memory=memory
        )

        sections[task.step] = content

    if progress_callback:
        progress_callback(total, total, "Completed")

    return sections    

# ======================================================
# DOCX Generator
# ======================================================

def create_docx(
    plan: ExecutionPlan,
    sections: Dict[int, str]
):

    os.makedirs("generated_docs", exist_ok=True)

    filename = os.path.join(
        "generated_docs",
        f"{safe_filename(plan.title)}_{uuid.uuid4().hex[:6]}.docx"
    )

    doc = Document()

    # ---------------------------
    # Title
    # ---------------------------

    title = doc.add_heading(plan.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()

    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle.add_run(
        f"Document Type: {plan.document_type}\n"
    ).bold = True

    subtitle.add_run(
        f"Generated on: {time.strftime('%d %B %Y')}"
    )

    doc.add_page_break()

    # ---------------------------
    # Sections
    # ---------------------------

    for task in plan.tasks:

        doc.add_heading(
            f"{task.step}. {task.description}",
            level=1
        )

        content = sections.get(task.step, "")

        paragraphs = content.split("\n")

        for p in paragraphs:

            p = p.strip()

            if not p:
                continue

            para = doc.add_paragraph()

            run = para.add_run(p)

            run.font.size = Pt(11)

        doc.add_paragraph()

    doc.save(filename)

    return filename


# ======================================================
# Main Agent
# ======================================================

def run_agent(
    user_request: str,
    progress_callback=None
):

    if len(user_request.strip()) < 10:

        raise Exception(
            "Request is too short."
        )

    # ---------------------------
    # Step 1
    # Autonomous Planning
    # ---------------------------

    plan = generate_plan(user_request)

    # ---------------------------
    # Step 2
    # Execute Tasks
    # ---------------------------

    sections = execute_plan(
        plan,
        user_request,
        progress_callback
    )

    # ---------------------------
    # Step 3
    # Build DOCX
    # ---------------------------

    filename = create_docx(
        plan,
        sections
    )

    # ---------------------------
    # Step 4
    # Return Result
    # ---------------------------

    return AgentResult(
        filename=filename,
        title=plan.title,
        document_type=plan.document_type,
        sections=sections
    )